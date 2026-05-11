from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)   # dark navy
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold_runs=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    # bold_runs: list of (substring, bold_flag) — not used here, kept simple
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Navy background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3964')
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = 'DDEEFF' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    doc.add_paragraph()   # spacing after table
    return table


# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(2)
tr = title.add_run('Mass Budget Assessment')
tr.font.name = 'Calibri'
tr.font.size = Pt(18)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(2)
sr = sub.add_run('OrbitGuard Assembly 3.2.2 — Thruster Configuration')
sr.font.name = 'Calibri'
sr.font.size = Pt(12)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

meta_lines = [
    ('Date',     datetime.date.today().strftime('%d %B %Y')),
    ('Prepared by', 'Karan Anand'),
    ('Status',   'Preliminary — Pending Material Verification'),
]
meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(4)
meta.paragraph_format.space_after  = Pt(12)
for label, val in meta_lines:
    run = meta.add_run(f'{label}:  ')
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run2 = meta.add_run(f'{val}     ')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)

doc.add_paragraph().add_run().add_break()   # thin divider

# ══════════════════════════════════════════════════════════════════════════════
# 1. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading('1.  Methodology')

body(
    'Part volumes were computed directly from the SolidWorks STEP assembly '
    '(OrbitGuard_Assembly_3.2.2_Thruster.STEP) using the OpenCASCADE geometry '
    'kernel. The assembly hierarchy was traversed via the XCAF framework to '
    'identify 67 unique part definitions across 320 total instances. Volume '
    'properties were evaluated on each closed B-rep solid; mass was then derived '
    'by applying material densities assigned per part on the basis of part '
    'nomenclature. COTS items and parts with non-descriptive identifiers were '
    'assigned a default aluminium density of 2700 kg/m³ pending confirmation. '
    'All results should be treated as preliminary until material assignments '
    'are verified against design records.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading('2.  Results')

body(
    'The estimated total assembly mass is 45.2 kg. Mass is heavily concentrated '
    'in the structural panels, which together account for approximately 75% of '
    'the total. The breakdown by category is given in Table 1.'
)

add_table(
    headers=['Category', 'Total Mass (g)', '% of Total'],
    rows=[
        ['Structural panels  (6 face panels + shear panel)', '33,882', '75.0%'],
        ['Propulsion  (2× BHT-350 T1 thrusters)',            ' 4,240', ' 9.4%'],
        ['Brackets and secondary structure',                  ' 3,965', ' 8.8%'],
        ['Electronics / RFIM assembly',                       '   752', ' 1.7%'],
        ['Fasteners  (bolts, washers)',                        '   116', ' 0.3%'],
        ['Other / unresolved COTS',                           '   230', ' 0.5%'],
        ['TOTAL',                                             '45,185', '100%'],
    ],
    col_widths=[9.5, 3.5, 3.0],
)

body(
    'Table 2 lists the ten highest-mass line items, sorted by total contribution. '
    'Panel geometry is modelled as solid aluminium; actual panel construction '
    '(likely sandwich with CFRP face sheets) would reduce the panel mass '
    'contribution significantly — see Section 3.'
)

add_table(
    headers=['Part', 'Qty', 'Vol/unit (cm³)', 'Density (kg/m³)', 'Mass/unit (g)', 'Total (g)'],
    rows=[
        ['Z− Panel',                      '1',  '5132',   '2700*', '13,857', '13,857'],
        ['Z+ Panel',                      '1',  '2081',   '2700*', ' 5,617', ' 5,617'],
        ['BHT-350 T1 thruster',           '2',  ' 785',   '2700*', ' 2,120', ' 4,240'],
        ['X+ Panel',                      '1',  '1154',   '2700*', ' 3,117', ' 3,117'],
        ['Y− Panel',                      '1',  '1117',   '2700*', ' 3,017', ' 3,017'],
        ['Y+ Panel',                      '1',  '1111',   '2700*', ' 3,001', ' 3,001'],
        ['X− Panel',                      '1',  '1110',   '2700*', ' 2,998', ' 2,998'],
        ['Shear Panel',                   '1',  ' 843',   '2700*', ' 2,275', ' 2,275'],
        ['Main Thruster Bracket',         '2',  ' 268',   '2700',  '   723', ' 1,446'],
        ['RS-00000821',                   '1',  ' 366',   '2700*', '   988', '   988'],
    ],
    col_widths=[5.5, 1.2, 3.0, 3.0, 3.0, 2.5],
)

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(6)
nr = note.add_run('* Density assumed from part name; pending material confirmation.')
nr.font.name = 'Calibri'
nr.font.size = Pt(9)
nr.font.italic = True
nr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# ══════════════════════════════════════════════════════════════════════════════
# 3. RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading('3.  Recommendations')

# R1
heading('3.1  Assign materials in SolidWorks', level=2)
body(
    'Material properties are not exported to STEP AP214, so the analysis above '
    'relies on density values inferred from part names. The most direct remedy '
    'is to assign the correct material to each part file in SolidWorks '
    '(Edit Material in the part document). Once materials are assigned, '
    'SolidWorks\' built-in Mass Properties tool (Evaluate → Mass Properties '
    'at assembly level) will report accurate masses directly from the model, '
    'without requiring any post-processing. This should be the primary reference '
    'for the mass budget going forward.'
)

# R2
heading('3.2  Confirm panel construction and material', level=2)
body(
    'The six face panels and the shear panel collectively represent 75% of the '
    'estimated assembly mass. All were modelled here as solid aluminium '
    '(2700 kg/m³). If the panels are sandwich construction with CFRP face '
    'sheets and aluminium honeycomb core — which is typical for this class of '
    'spacecraft — the effective density is in the range of 600–900 kg/m³ for '
    'the panel as a whole. This correction alone could reduce the total mass '
    'estimate by 8–12 kg. The mechanical team should confirm the panel material '
    'specification and, if the CAD models are simplified solids rather than '
    'true sandwich representations, override the panel densities accordingly '
    'in the mass budget tool.'
)

# R3
heading('3.3  Verify COTS component masses against datasheets', level=2)
body(
    'Geometric volume is not a reliable mass estimator for COTS components, '
    'which are typically represented in assemblies as simplified or defeatured '
    'bodies. The following items in particular should be replaced with '
    'datasheet-quoted masses:'
)

items = [
    ('BHT-350 T1 (Busek Hall-effect thruster)',
     'The script returns 2.12 kg per unit from a geometry that may be a '
     'placeholder model. The Busek datasheet should be the authoritative source.'),
    ('RS-00000821, RS-00000873, RS-00000450',
     'Part numbers not resolved from the assembly name. Identify via the BOM '
     'and source from the supplier datasheet.'),
    ('COTS-000093, COTS-000110, COTS-000111, COTS-000112, COTS-000114',
     'Internal COTS designators. Cross-reference with the procurement register '
     'to obtain quoted masses.'),
]
for title_text, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    r1 = p.add_run(title_text + ': ')
    r1.font.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)

doc.add_paragraph()

# R4
heading('3.4  Address duplicated part names in STEP export', level=2)
body(
    'SolidWorks AP214 exports are appending each part\'s name to itself '
    '(e.g. "Z-_Panel_-(2)_Z-_Panel_-(2)"). While this does not affect the '
    'volume computation, it complicates BOM readability and downstream tooling. '
    'The issue can be resolved by selecting AP242 format on export, or by '
    'adjusting the STEP export options under File → Save As → Options '
    '(uncheck "Export IDs").'
)

# R5
heading('3.5  Establish a controlled mass budget baseline', level=2)
body(
    'Once material assignments and COTS masses are confirmed, a formal mass '
    'budget baseline should be established with mass margins applied per the '
    'project\'s systems engineering margin policy (typically 20% at component '
    'level at this stage). The mass budget tool used here can be re-run against '
    'any updated STEP export and accepts a JSON override file for COTS and '
    'material corrections, allowing iterative updates without re-parsing the '
    'full assembly each time from scratch.'
)

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
out = '/home/karan.anand/Documents/PythonScripts/OrbitGuard_3.2.2_Mass_Budget_Assessment.docx'
doc.save(out)
print(f'Saved: {out}')
