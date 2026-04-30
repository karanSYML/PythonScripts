#!/usr/bin/env python3
"""
build_csv_guide.py
==================
Generates the external consultant CSV data package guide as a .docx file.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── Colour palette ─────────────────────────────────────────────────────────────
DARK     = RGBColor(0x1E, 0x29, 0x3B)   # slate-900
MID      = RGBColor(0x33, 0x41, 0x55)   # slate-700
LIGHT    = RGBColor(0x64, 0x74, 0x8B)   # slate-500
ACCENT   = RGBColor(0x1D, 0x4E, 0xD8)   # blue-700
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
HDR_BG   = "1E293B"   # table header fill (hex string for XML)
ALT_BG   = "F1F5F9"   # alternating row fill
BAND_BG  = "E2E8F0"   # section band


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CBD5E1", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def style_header_row(row, font_size=9):
    for cell in row.cells:
        set_cell_bg(cell, HDR_BG)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold      = True
                run.font.size      = Pt(font_size)
                run.font.name      = "Calibri"


def style_data_row(row, alt=False, font_size=9):
    bg = ALT_BG if alt else "FFFFFF"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        set_cell_borders(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                run.font.size      = Pt(font_size)
                run.font.name      = "Calibri"


def add_table(doc, headers, rows, col_widths_cm, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, (cell, h) in enumerate(zip(hdr.cells, headers)):
        cell.text = h
        cell.width = Cm(col_widths_cm[i])
    style_header_row(hdr, font_size)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.text  = str(val)
            cell.width = Cm(col_widths_cm[i])
        style_data_row(row, alt=(r_idx % 2 == 1), font_size=font_size)

    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style    = "Normal"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.bold  = True
    run.font.size  = Pt(13 if level == 1 else 11)
    run.font.color.rgb = DARK if level == 1 else ACCENT
    return p


def add_body(doc, text, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name      = "Calibri"
    run.font.size      = Pt(10)
    run.font.color.rgb = MID
    run.font.italic    = italic
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name      = "Calibri Light"
    run.font.size      = Pt(9)
    run.font.color.rgb = LIGHT
    run.font.italic    = True
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name      = "Courier New"
    run.font.size      = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return p


def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after  = Pt(2)
    title.paragraph_format.space_before = Pt(0)
    tr = title.add_run("CSV Data Package — Structure and Content Guide")
    tr.font.name      = "Calibri"
    tr.font.bold      = True
    tr.font.size      = Pt(18)
    tr.font.color.rgb = DARK

    meta_lines = [
        ("Prepared for:", "External Consultant"),
        ("Dataset:",       "Endurance FR RDV — GEO Rendezvous Mission (Simulation Run, Sep 2028)"),
        ("Prepared by:",   "Karan Anand"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        lb = p.add_run(f"{label}  ")
        lb.font.name      = "Calibri"
        lb.font.bold      = True
        lb.font.size      = Pt(9)
        lb.font.color.rgb = LIGHT
        vr = p.add_run(value)
        vr.font.name      = "Calibri"
        vr.font.size      = Pt(9)
        vr.font.color.rgb = MID

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section 1: Files provided ─────────────────────────────────────────────
    add_heading(doc, "1.  Files Provided")
    add_body(doc,
        "Four CSV files are provided, one per combination of simulation run and attitude mode. "
        "Each file covers the full mission timeline — rendezvous approach (RDV) followed directly "
        "by the inspection phase (INS) — at a uniform propagation timestep of 300 seconds.")

    add_table(doc,
        headers=["File", "Attitude Mode", "Rows", "Columns"],
        rows=[
            ["og3_target_SunOpt_nominal.csv", "Mode 1 — Target+Sun  (og3 run)", "6,691", "40"],
            ["og3_Nadir_sunOpt_nominal.csv",  "Mode 2 — Nadir+Sun   (og3 run)", "6,691", "40"],
            ["end1_target_sunOpt.csv",         "Mode 1 — Target+Sun  (end1 run)", "6,621", "49"],
            ["end1_nadir_sunopt.csv",           "Mode 2 — Nadir+Sun   (end1 run)", "6,621", "49"],
        ],
        col_widths_cm=[6.5, 6.5, 2.0, 2.0],
    )

    add_body(doc,
        "The column count differs between og3 (40 columns) and end1 (49 columns) because the two "
        "runs use propulsion models with different fidelity: og3 models 8 RCS thruster channels "
        "and 2 PPS firing values; end1 models 16 and 3 respectively. The scenario, trajectory, "
        "and attitude data are structured identically in both runs.", space_after=4)

    add_note(doc,
        "Mode 1 (Target+Sun): primary boresight (+Z body axis) constrained to the rendezvous target; "
        "solar array normal optimised toward the Sun.\n"
        "Mode 2 (Nadir+Sun): primary boresight (+Z body axis) constrained to the local nadir "
        "(Earth centre direction); solar array normal optimised toward the Sun.")

    # ── Section 2: Overall structure ──────────────────────────────────────────
    add_heading(doc, "2.  Overall Structure")
    add_body(doc,
        "Each file is a flat, single-header CSV. Every row corresponds to one propagation timestep "
        "and is indexed by a UTC timestamp (column UTC_ISOC). The file is continuous — no separator "
        "rows or gaps exist between the RDV and INS phases. The phase boundary can be identified "
        "programmatically from the along-track separation column a_dlambda_m: it occurs at the "
        "first epoch where this value crosses −5,000 m (approximately mission day 15.1 for end1, "
        "day 15.2 for og3).")
    add_body(doc,
        "The columns fall into two logical groups, both present in every row:")

    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run("Group 1 — State Data  ")
    r1.font.name = "Calibri"; r1.font.bold = True
    r1.font.size = Pt(10); r1.font.color.rgb = DARK
    r2 = p.add_run("(columns 1–28):  trajectory, attitude, and pointing quantities at each timestep. "
                   "All rows are fully populated.")
    r2.font.name = "Calibri"; r2.font.size = Pt(10); r2.font.color.rgb = MID

    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run("Group 2 — Propulsion Event Data  ")
    r1.font.name = "Calibri"; r1.font.bold = True
    r1.font.size = Pt(10); r1.font.color.rgb = DARK
    r2 = p.add_run("(columns 29 onward):  firing data extracted from the MATLAB simulation packages "
                   "and in-filled into the matching time-series row. All firing columns are zero "
                   "for rows with no propulsion event — there are no blank or NaN cells in the file.")
    r2.font.name = "Calibri"; r2.font.size = Pt(10); r2.font.color.rgb = MID

    # ── Section 3: Column reference ───────────────────────────────────────────
    add_heading(doc, "3.  Column Reference")

    add_heading(doc, "3.1   Group 1 — State Data (columns 1–28, all rows populated)", level=2)

    add_table(doc,
        headers=["Column name(s)", "Description", "Units"],
        rows=[
            ["UTC_ISOC",
             "Timestamp in ISO 8601 format, UTC, nanosecond precision.",
             "—"],
            ["pos_Earth2satCoG_ECI_m_x/y/z",
             "Servicer centre-of-gravity position vector in the Earth-Centred Inertial (ECI) frame.",
             "m"],
            ["vel_Earth2satCoG_ECI_ms_x/y/z",
             "Servicer velocity vector in the ECI frame.",
             "m/s"],
            ["pos_Earth2tgtCoG_ECI_m_x/y/z",
             "Rendezvous target centre-of-gravity position vector in the ECI frame.",
             "m"],
            ["pos_Earth2Sun_ECI_m_x/y/z",
             "Sun position vector in the ECI frame.",
             "m"],
            ["quat_ECI2MRF_a/i/j/k",
             "Attitude quaternion transforming vectors from ECI to the Mission Reference Frame (body frame). "
             "Scalar-first convention: a = w (scalar part), i/j/k = x/y/z (vector part).",
             "—"],
            ["angle_SunPhaseAngle_rad",
             "Sun phase angle at the current epoch.",
             "rad"],
            ["uv_sat2Earth_MRF_x/y/z",
             "Unit vector from the spacecraft to the Earth centre, expressed in the body frame. "
             "Pre-computed pointing reference; consistent with the attitude quaternion.",
             "—"],
            ["uv_sat2Target_MRF_x/y/z",
             "Unit vector from the spacecraft to the rendezvous target, expressed in the body frame.",
             "—"],
            ["angle_sunFromMRFaxes_deg_x/y/z",
             "Angle between the Sun direction vector and each body axis: "
             "x = angle from +X axis, y = angle from +Y axis, z = angle from +Z axis (camera boresight). "
             "Used directly for thermal constraint analysis.",
             "deg"],
            ["a_dlambda_m",
             "Along-track relative separation between servicer and target, computed along the GEO arc "
             "(servicer minus target). Negative values indicate the servicer is trailing the target. "
             "Phase boundary occurs at −5,000 m.",
             "m"],
        ],
        col_widths_cm=[5.5, 9.0, 1.5],
        font_size=9,
    )

    add_heading(doc, "3.2   Group 2 — Propulsion Event Data (columns 29 onward)", level=2)
    add_body(doc,
        "These columns are populated only for rows that coincide with a recorded propulsion event. "
        "All other rows carry the value 0 across every firing column. The state data columns "
        "(Group 1) remain fully populated on firing rows — both datasets are always present together.")

    add_table(doc,
        headers=["Column name(s)", "Description", "Units"],
        rows=[
            ["maneuver_type",
             "Integer event flag identifying the propulsion system active at this timestep.\n"
             "  0 = no firing event\n"
             "  1 = PPS firing (Primary Propulsion System — main engine)\n"
             "  2 = RCS firing (Reaction Control System)",
             "integer"],
            ["manTime_hr",
             "Mission elapsed time of the firing event as recorded in the simulation data package. "
             "This is the exact firing epoch, not the rounded grid timestamp — see Section 5 for the "
             "distinction.",
             "hours"],
            ["firingTime_rcs_T01 … T08 (og3)\nfiringTime_rcs_T01 … T16 (end1)",
             "Per-thruster firing duration for each RCS thruster channel. "
             "The number of channels reflects the propulsion model fidelity of each simulation run. "
             "Zero on all non-RCS rows.",
             "s"],
            ["firingTime_pps_1, _2 (og3)\nfiringTime_pps_1, _2, _3 (end1)",
             "Firing values for each PPS (main engine) channel. "
             "Zero on all non-PPS rows.",
             "s"],
        ],
        col_widths_cm=[5.5, 9.0, 1.5],
        font_size=9,
    )

    # ── Section 4: Propulsion system distinction ──────────────────────────────
    add_heading(doc, "4.  Propulsion System Distinction")

    add_body(doc,
        "PPS — Primary Propulsion System (main engine).  "
        "Large impulsive corrections applied exclusively during the far-range rendezvous phase. "
        "Both simulation runs record 80 PPS events, all occurring before the phase boundary "
        "(i.e. within the RDV portion of the timeline). No PPS activity is recorded during the "
        "inspection phase.")
    add_body(doc,
        "RCS — Reaction Control System.  "
        "Discrete attitude control and trajectory correction pulses distributed across both phases: "
        "20 events during the far-range RDV phase and 24 during the inspection phase (44 total). "
        "RCS firings are present in both phases across both simulation runs.")

    add_table(doc,
        headers=["", "PPS (maneuver_type = 1)", "RCS (maneuver_type = 2)"],
        rows=[
            ["Mission phase",      "RDV only",                    "RDV + INS"],
            ["Events (both runs)", "80",                           "44  (20 RDV + 24 INS)"],
            ["Channels — og3",     "2  (firingTime_pps_1…2)",     "8   (firingTime_rcs_T01…T08)"],
            ["Channels — end1",    "3  (firingTime_pps_1…3)",     "16  (firingTime_rcs_T01…T16)"],
            ["Zero-firing records","1 event (PPS, null burst)",   "6 events (null bursts)"],
        ],
        col_widths_cm=[4.5, 5.5, 6.0],
        font_size=9,
    )

    add_note(doc,
        "Note: a small number of recorded firing events carry all-zero firing durations "
        "(1 PPS event and 6 RCS events across both datasets). These are genuine entries in "
        "the simulation record — the epoch was logged but no thrust was executed. They are "
        "identifiable by maneuver_type ≠ 0 combined with all firingTime columns equal to 0.")

    # ── Section 5: Timestamp alignment ───────────────────────────────────────
    add_heading(doc, "5.  Timestamp Alignment")
    add_body(doc,
        "Firing event epochs are stored in the MATLAB simulation data packages as absolute seconds "
        "from the J2000 epoch (2000-01-01T12:00:00 UTC). These epochs do not fall exactly on the "
        "300-second propagation grid. Each firing event has been matched to the nearest propagation "
        "timestep using a minimum time-difference criterion. The worst-case offset between the "
        "recorded firing epoch and the UTC_ISOC timestamp of its assigned row is 148.9 seconds "
        "— well within the 150-second half-step tolerance.")
    add_body(doc,
        "The manTime_hr column always preserves the exact firing epoch from the simulation record "
        "expressed as mission elapsed hours. It is therefore not derived from UTC_ISOC and will "
        "not be exactly consistent with it — this is expected and by design.")

    # ── Section 6: Filtering ──────────────────────────────────────────────────
    add_heading(doc, "6.  Filtering and Usage")
    add_body(doc,
        "The maneuver_type column is the primary filter for separating the time-series and event "
        "subsets. The following Python example illustrates standard usage:")

    for line in [
        "import pandas as pd",
        "",
        "df = pd.read_csv('og3_target_SunOpt_nominal.csv')",
        "",
        "# Continuous time-series (attitude, trajectory, pointing)",
        "ts   = df[df['maneuver_type'] == 0]   # 6,567 rows",
        "",
        "# PPS firing events only",
        "pps  = df[df['maneuver_type'] == 1]   # 80 rows",
        "",
        "# RCS firing events only",
        "rcs  = df[df['maneuver_type'] == 2]   # 44 rows",
        "",
        "# Phase boundary (RDV → INS)",
        "ins_start = df[df['a_dlambda_m'] > -5000].index[0]",
    ]:
        add_code(doc, line)

    add_body(doc,
        "All firing columns are numeric (float64). No NaN values are present anywhere in the file. "
        "Empty cells do not exist — every cell is either a real value or zero.")

    # ── Section 7: What is not included ───────────────────────────────────────
    add_heading(doc, "7.  What Is Not Included")
    add_body(doc,
        "These files contain the simulation state history and propulsion event record only. "
        "Post-processed analysis outputs — angular separation between Earth and target, eigenaxis "
        "slew angles, thermal constraint metrics (camera sun exclusion, star tracker blinding), "
        "and antenna feasibility assessments — are not embedded in the CSV. They are produced "
        "separately by the Python analysis suite and are available as figures and printed summaries "
        "on request.")

    out = "CSV_Data_Package_Guide.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
