#!/usr/bin/env python3
"""
update_document.py
==================
Updates EarthTargetPointingFeasibility_OG3.docx with:
  - Mission sub-phases table added to Section 1
  - Section 3.2 / 3.3 headings updated with phase names
  - Figures inserted throughout
  - New Section 5.3 for Telescope (+Z, 3.75° sun exclusion)
  - Old 5.3 Mode guidance renumbered to 5.4
  - New Section 6 Power Budget
  - Conclusion renumbered to 7, Open Items to 8
"""

import os
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = "EarthTargetPointingFeasibility_OG3.docx"
DEST = "EarthTargetPointingFeasibility_OG3_v2.docx"

IMG = {
    "fig1a":         "og3_fig1a_earth_target_far.png",
    "fig1b":         "og3_fig1b_earth_target_close.png",
    "feasibility_far":   "og3_feasibility_far.png",
    "feasibility_close": "og3_feasibility_close.png",
    "slew":          "og3_slew_comparison.png",
    "thermal_far":   "og3_thermal_far.png",
    "thermal_close": "og3_thermal_close.png",
    "power_far":     "power_m1_far.png",
    "power_close":   "power_m1_close.png",
}

PHASES = [
    ("Phase 1", "Station-keeping",       "1 day",   "−60 km"),
    ("Phase 2", "Approach",              "10 days", "−60 km → −30 km"),
    ("Phase 3", "Station-keeping",       "1 day",   "−30 km"),
    ("Phase 4", "Accelerated approach",  "4 days",  "−30 km → −1 km"),
    ("Phase 5", "Station-keeping",       "1 day",   "−1 km"),
    ("Phase 6", "Checkpoint resizing",   "2 days",  "−1 km"),
    ("Phase 7", "Station-keeping",       "1 day",   "−1 km"),
    ("Phase 8", "Fly-by",                "2 days",  "−1 km → +1 km"),
    ("Phase 9", "Station-keeping",       "1 day",   "+1 km"),
]


# ── XML helpers ───────────────────────────────────────────────────────────────

def _move_after(ref_elem, new_elem):
    """Detach new_elem from wherever it is and place it after ref_elem."""
    parent = new_elem.getparent()
    if parent is not None:
        parent.remove(new_elem)
    ref_elem.addnext(new_elem)


def _find_para(doc, text_fragment):
    """Return first paragraph whose text contains text_fragment."""
    for p in doc.paragraphs:
        if text_fragment in p.text:
            return p
    return None


def _find_all_paras(doc, text_fragment):
    """Return all paragraphs whose text contains text_fragment."""
    return [p for p in doc.paragraphs if text_fragment in p.text]


def _add_centered_picture(doc, ref_elem, image_path, width_in, caption_text):
    """Insert a centered picture + caption paragraph pair after ref_elem."""
    if not os.path.exists(image_path):
        print(f"  WARNING: image not found: {image_path}")
        return ref_elem

    # -- picture paragraph --
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    _move_after(ref_elem, pic_para._element)

    # -- caption paragraph --
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption_text)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    _move_after(pic_para._element, cap_para._element)

    return cap_para._element   # return tail element for chaining


def _add_heading_after(doc, ref_elem, text, level):
    """Insert a heading paragraph after ref_elem."""
    h = doc.add_heading(text, level=level)
    _move_after(ref_elem, h._element)
    return h._element


def _add_para_after(doc, ref_elem, text, style="Normal", bold=False):
    """Insert a Normal paragraph after ref_elem."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    _move_after(ref_elem, p._element)
    return p._element


def _add_bullet_after(doc, ref_elem, text):
    """Insert a List Bullet paragraph after ref_elem."""
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    _move_after(ref_elem, p._element)
    return p._element


def _bulk_replace_runs(doc, mapping):
    """Walk every paragraph (incl. table cells) and apply text replacement at the
    run level. Replacement is only applied when the search substring lies within
    a single run (sufficient for this document's flat formatting)."""
    def _apply(paragraph):
        for run in paragraph.runs:
            for old, new in mapping.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    for p in doc.paragraphs:
        _apply(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _apply(p)


def _add_table_after(doc, ref_elem, headers, rows):
    """Insert a table with headers + rows after ref_elem."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    _move_after(ref_elem, tbl._element)
    return tbl._element


# ── Main document update ──────────────────────────────────────────────────────

def update():
    doc = Document(SRC)

    # ── 1. Add mission sub-phases table to Section 1 ─────────────────────────
    print("  Adding mission sub-phases table...")
    # Insert after the last list bullet in section 1 (antenna entry)
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    last_bullet = bullet_paras[-1] if bullet_paras else _find_para(doc, "Camera and antenna")

    # Intro sentence
    ref = _add_para_after(
        doc, last_bullet._element,
        "The full rendezvous timeline (−60 km to +1 km) is structured into nine "
        "mission sub-phases as summarised below. All analyses in this document are "
        "presented against this sub-phase timeline.")

    # Table
    headers = ["Phase", "Name", "Duration", "Along-track position"]
    rows    = [(p[0], p[1], p[2], p[3]) for p in PHASES]
    ref = _add_table_after(doc, ref, headers, rows)

    # ── 2. Update Section 3.2 heading ────────────────────────────────────────
    print("  Updating section 3.2 heading...")
    h32 = _find_para(doc, "3.2 Far-range phase")
    if h32:
        for run in h32.runs:
            run.text = ""
        h32.runs[0].text = (
            "3.2 Phases 1–4 (Phase 1 – Station-keeping through "
            "Phase 4 – Accelerated approach, days 0–15.2, −60 to −5 km)")

    # ── 3. Update Section 3.3 heading ────────────────────────────────────────
    print("  Updating section 3.3 heading...")
    h33 = _find_para(doc, "3.3 Close-range phase")
    if h33:
        for run in h33.runs:
            run.text = ""
        h33.runs[0].text = (
            "3.3 Phases 5–9 (Phase 5 – Station-keeping through "
            "Phase 9 – Station-keeping, days 15.2–23.2, −5 to +1 km)")

    # ── 4. Rename 5.3 → 5.4 ──────────────────────────────────────────────────
    print("  Renumbering 5.3 → 5.4...")
    h53 = _find_para(doc, "5.3 Mode selection")
    if h53:
        old = h53.runs[0].text if h53.runs else h53.text
        new = old.replace("5.3", "5.4")
        if h53.runs:
            h53.runs[0].text = new

    # ── 5. Rename 6. Conclusion → 7. Conclusion ──────────────────────────────
    print("  Renumbering Conclusion and Open Items...")
    h6 = _find_para(doc, "6. Conclusion")
    if h6 and h6.runs:
        h6.runs[0].text = h6.runs[0].text.replace("6.", "7.")
    h7 = _find_para(doc, "7. Open Items")
    if h7 and h7.runs:
        h7.runs[0].text = h7.runs[0].text.replace("7.", "8.")

    # ── 5b. Bulk replace "far/close range" terminology in remaining prose ────
    print("  Replacing far/close range terminology with Phase nomenclature...")
    # Order matters: longer/specific phrases first to preserve grammar.
    mapping = {
        # Specific phrasings that need verb-agreement or article fixes
        "During the far-range rendezvous":     "During the rendezvous",
        "The far-range phase is":              "Phases 1–4 are",
        "The close-range phase is":            "Phases 5–9 are",
        "The close-range phase presents":      "Phases 5–9 present",
        "The close-range phase delivers":      "Phases 5–9 deliver",
        "the far-range phase delivers":        "Phases 1–4 deliver",
        "the close-range phase delivers":      "Phases 5–9 deliver",
        # Common prepositional forms
        "in the far-range phase":              "across Phases 1–4",
        "in the close-range phase":            "across Phases 5–9",
        "during the far-range phase":          "across Phases 1–4",
        "during the close-range phase":        "across Phases 5–9",
        "during the far-range approach":       "during the Phases 1–4 approach",
        "during the close-range approach":     "during the Phases 5–9 approach",
        "compared to the far-range phase":     "compared to Phases 1–4",
        "compared to the close-range phase":   "compared to Phases 5–9",
        "at far range":                        "across Phases 1–4",
        "at close range":                      "across Phases 5–9",
        "in the far range":                    "across Phases 1–4",
        "in the close range":                  "across Phases 5–9",
        # Generic forms (catch-alls)
        "the far-range phase":   "Phases 1–4",
        "the close-range phase": "Phases 5–9",
        "Far-range phase":       "Phases 1–4",
        "Close-range phase":     "Phases 5–9",
        "far-range phase":       "Phases 1–4",
        "close-range phase":     "Phases 5–9",
        "Far-range":             "Phases 1–4",
        "Close-range":           "Phases 5–9",
        "far-range":             "Phases 1–4",
        "close-range":           "Phases 5–9",
        "Far range":             "Phases 1–4",
        "Close range":           "Phases 5–9",
        " far range":            " Phases 1–4",
        " close range":          " Phases 5–9",
        "far range":             "Phases 1–4",
        "close range":           "Phases 5–9",
    }
    _bulk_replace_runs(doc, mapping)

    # ── 6. Insert figures after Section 2 ────────────────────────────────────
    print("  Inserting Section 2 figures...")
    sec2_last = _find_para(doc, "An important property of the 90°")
    if sec2_last:
        ref = sec2_last._element
        ref = _add_centered_picture(doc, ref, IMG["fig1a"], 6.2,
            "Figure 1a — Earth–Target angular separation, Phases 1–4 "
            "(−60 to −5 km, Station-keeping through Accelerated approach). "
            "Separation is tightly bounded near 90° throughout.")
        ref = _add_centered_picture(doc, ref, IMG["fig1b"], 6.2,
            "Figure 1b — Earth–Target angular separation, Phases 5–9 "
            "(−5 to +1 km, Station-keeping through Fly-by). Large oscillations "
            "with 90° crossings every ≈11.7 h (mean half-period).")

    # ── 7. Insert figures after Section 3.2 ──────────────────────────────────
    print("  Inserting Section 3.2 figure...")
    sec32_last = _find_para(doc, "Propulsion context: 72 PPS")
    if sec32_last:
        ref = _add_centered_picture(doc, sec32_last._element, IMG["feasibility_far"], 6.2,
            "Figure 2 — Pointing feasibility, Phases 1–4 (−60 to −5 km). "
            "Mode 1 antenna error (red) remains within the X-band 4.5° half-cone "
            "for 92.7% of epochs. Mode 2 achieves 0% feasibility.")

    # ── 8. Insert figures after Section 3.3 ──────────────────────────────────
    print("  Inserting Section 3.3 figure...")
    sec33_last = _find_para(doc, "RCS manoeuvre context: 44 RCS")
    if sec33_last:
        ref = _add_centered_picture(doc, sec33_last._element, IMG["feasibility_close"], 6.2,
            "Figure 3 — Pointing feasibility, Phases 5–9 (−5 to +1 km). "
            "Feasibility drops to 15.5% within the X-band cone; the oscillation "
            "period of ≈23.4 h exceeds the 6-hour ConOps window interval.")

    # ── 9. Insert figure after Section 4 ─────────────────────────────────────
    print("  Inserting Section 4 figure...")
    sec4_last = _find_para(doc, "ConOps implication: if the spacecraft is in Mode 1")
    if sec4_last:
        ref = _add_centered_picture(doc, sec4_last._element, IMG["slew"], 6.5,
            "Figure 4 — Eigenaxis slew cost from Mode 1 (blue) and Mode 2 (red) "
            "into the combined Earth+Target attitude. Left: Phases 1–4 (−60 to −5 km); "
            "Right: Phases 5–9 (−5 to +1 km). Mode 1 is cheaper for 57.1% of epochs.")

    # ── 10. Insert figures after Section 5.1 ─────────────────────────────────
    print("  Inserting Section 5.1 figures...")
    sec51_last = _find_para(doc, "consistently high violation rate (~33%)")
    if sec51_last:
        ref = sec51_last._element
        ref = _add_centered_picture(doc, ref, IMG["thermal_far"], 6.2,
            "Figure 5a — Thermal constraints, Phases 1–4 (−60 to −5 km). "
            "Top: Camera/NAC (+Z, 30° exclusion). Middle: Telescope (+Z, 3.75° exclusion). "
            "Bottom: STR (+X, 35° exclusion). Amber/red spans mark ≤5 min / >5 min "
            "Telescope violations respectively.")
        ref = _add_centered_picture(doc, ref, IMG["thermal_close"], 6.2,
            "Figure 5b — Thermal constraints, Phases 5–9 (−5 to +1 km). "
            "Mode 1 Telescope violations drop to zero across Phases 5–9.")

    # ── 11. Add Telescope section 5.3 before old 5.3 (now 5.4) ──────────────
    print("  Adding Telescope section 5.3...")
    # Insert before the (now renumbered) 5.4 heading
    h54 = _find_para(doc, "5.4 Mode selection")
    if h54:
        # Work backwards: insert new section BEFORE h54, i.e. after the STR last para
        str_last = _find_para(doc, "STR sun blinding is negligible in both modes")
        if str_last:
            ref = str_last._element
        else:
            ref = h54._element.getprevious()

        ref = _add_heading_after(doc, ref,
            "5.3 Telescope sun exclusion (+Z axis, 3.75° half-cone)", level=2)

        ref = _add_para_after(doc, ref,
            "The Telescope replaces the WAC and is co-boresighted with the NAC on the "
            "+Z body axis. Its sun exclusion half-cone is 3.75° with a hard operational "
            "constraint of no more than 5 minutes of continuous solar exposure within "
            "the exclusion cone. Telescope power consumption matches the NAC (5 W idle, "
            "8 W active).")

        ref = _add_para_after(doc, ref,
            "Phases 1–4 (−60 to −5 km, Station-keeping through Accelerated approach): "
            "Mode 1 records 1.98% violation epochs with 15 continuous events exceeding "
            "the 5-minute limit, the longest lasting 35 minutes — approximately 7× the "
            "allowed maximum. Mode 2 is worse, with 4.10% violation and 30 events "
            "exceeding 5 minutes (longest: 30 minutes). These violations recur at the "
            "GEO synodic period as the +Z boresight sweeps through the solar exclusion "
            "cone.")

        ref = _add_para_after(doc, ref,
            "Phases 5–9 (−5 to +1 km, Station-keeping through Fly-by and final "
            "Station-keeping): Mode 1 records zero violations. As the along-track "
            "separation closes, the +Z target boresight progressively separates from "
            "the Sun direction, providing complete Telescope thermal compliance. Mode 2 "
            "retains 4.21% violation with 16 events exceeding 5 minutes (longest: "
            "30 minutes).")

        ref = _add_para_after(doc, ref,
            "Operational implication: Mode 1 across Phases 5–9 is fully compliant with "
            "the Telescope sun exclusion constraint. Across Phases 1–4, Mode 1 produces "
            "intermittent violations driven by the GEO synodic geometry. The 35-minute "
            "worst-case event should be reviewed by the instrument team against the "
            "Telescope thermal qualification envelope. Operator scheduling of observation "
            "windows to avoid known violation epochs in Phase 2 – Approach and "
            "Phase 4 – Accelerated approach is recommended.")

    # ── 12. Add Power Budget section 6 before old Conclusion (now 7) ─────────
    print("  Adding Power Budget section 6...")
    h7_conclusion = _find_para(doc, "7. Conclusion")
    if h7_conclusion:
        ref = h7_conclusion._element.getprevious()

        ref = _add_heading_after(doc, ref, "6. Power Budget Analysis", level=1)

        ref = _add_para_after(doc, ref,
            "Solar power generation and subsystem consumption are evaluated across the "
            "full OG3 nominal mission timeline. Generation is modelled as "
            "P_gen = 344 × sin(θ_Y) W (EOL), where θ_Y is the instantaneous angle "
            "between the Sun and the +Y body axis; the solar array drive (SAD) maintains "
            "near-optimal sun-pointing within the body X-Z plane at all times. Platform "
            "power constants are drawn from the OG3 system database.")

        ref = _add_heading_after(doc, ref, "6.1 Solar generation and eclipse", level=2)

        ref = _add_para_after(doc, ref,
            "Generation follows the 23.4-hour GEO synodic oscillation. Significant "
            "eclipse periods occur across all phases, consistent with the mission epoch "
            "near the September 2028 equinox (308 eclipse epochs at 300-second timestep "
            "resolution over the 23-day mission). Mode 1 (Target+Sun) achieves a mean "
            "solar generation of 199 W over the full timeline. The secondary "
            "Sun-optimisation constraint keeps the SAD well-oriented in most epochs.")

        ref = _add_heading_after(doc, ref, "6.2 Power balance by phase", level=2)

        ref = _add_para_after(doc, ref,
            "Phases 1–4 (−60 to −5 km, days 0–15.2): mean generation 218 W against "
            "mean consumption 285 W (Mode 1, baseline 260 W + 16 W Camera+Telescope "
            "in-window / 10 W out-of-window), yielding a mean deficit of approximately "
            "−66 W. Surplus generation occurs in 40% of epochs. PPS firings (72 events "
            "across Phases 1–4) drive the bus to ≈640 W (260 W baseline + 380 W PPU "
            "average draw, per POELIS pps.average_power_consumption_max), producing "
            "worst-case deficits of ≈−650 W when a firing coincides with eclipse. "
            "Battery cycling is continuous and, in this nominal scenario, the integrated "
            "state-of-charge drops to the 20% DoD floor during extended eclipse-plus-"
            "firing combinations.")

        ref = _add_para_after(doc, ref,
            "Phases 5–9 (−5 to +1 km, days 15.2–23.2): mean generation drops to 162 W "
            "as eclipse duration increases near equinox; mean consumption is 271 W, "
            "yielding a mean deficit of −110 W. Surplus generation occurs in only 21% "
            "of epochs. Battery state-of-charge stays at or near the 20% DoD floor "
            "through much of Phase 8 – Fly-by and Phase 9 – Station-keeping. RCS "
            "manoeuvres (44 events, exclusive to Phases 5–9) contribute short "
            "consumption spikes that are small relative to eclipse-driven deficits.")

        ref = _add_heading_after(doc, ref, "6.3 ConOps implications", level=2)

        bullets = [
            "PPS firings should be scheduled during eclipse-free periods where possible "
            "to avoid compounding solar deficit with peak bus consumption "
            "(≈640 W: 260 W baseline + 380 W PPU average draw).",
            "The 550 Wh battery provides approximately 2 hours of autonomy at the "
            "260 W baseline under full eclipse conditions.",
            "X-band transmitter power (40 W) and Telescope power (5 W idle, 8 W active, "
            "matching NAC) are applied during combined pointing windows. WAC has been "
            "removed from the payload configuration.",
            "Phase 8 – Fly-by and Phase 9 – Station-keeping represent the tightest "
            "power margins of the mission. Battery management should be explicitly "
            "included in the Phase 8 ConOps.",
        ]
        for b in bullets:
            ref = _add_bullet_after(doc, ref, b)

        # Figures
        ref = _add_centered_picture(doc, ref, IMG["power_far"], 6.2,
            "Figure 6a — Power budget, Phases 1–4 (−60 to −5 km), Mode 1 (Target+Sun). "
            "Top: solar generation with eclipse shading. Second: consumption stacked by "
            "subsystem. Third: net power balance (green = surplus, red = deficit). "
            "Bottom: battery state-of-charge.")

        ref = _add_centered_picture(doc, ref, IMG["power_close"], 6.2,
            "Figure 6b — Power budget, Phases 5–9 (−5 to +1 km), Mode 1. "
            "Generation decreases with increased eclipse depth. Battery SoC approaches "
            "the 20% DoD limit through Phases 8–9.")

    doc.save(DEST)
    print(f"\n  Saved: {DEST}")


if __name__ == "__main__":
    print(f"\n{'='*60}")
    print("  Document Update")
    print(f"{'='*60}\n")
    update()
    print("\nDone.")
